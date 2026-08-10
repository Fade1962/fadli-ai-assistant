from docx import Document

def create_docx(filename, content):
    doc = Document()
    doc.add_heading("NARA", level=1)
    for line in (content or "").splitlines():
        doc.add_paragraph(line)
    doc.save(filename)
    return filename
