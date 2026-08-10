from docx import Document


def process_document(filename):

    try:

        document = Document(
            filename
        )

        paragraphs = []


        # ==============================
        # PARAGRAPH
        # ==============================

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                paragraphs.append(
                    text
                )


        # ==============================
        # TABLE
        # ==============================

        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    value = cell.text.strip()

                    cells.append(value)


                if any(cells):

                    paragraphs.append(
                        " | ".join(cells)
                    )


        result = "\n\n".join(
            paragraphs
        ).strip()


        if not result:

            return (
                "DOCX berhasil dibaca, "
                "tetapi tidak ditemukan teks."
            )


        return result


    except Exception as error:

        print(
            "DOCUMENT PROCESS ERROR:",
            repr(error)
        )

        raise
